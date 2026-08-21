from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


def find_pandoc() -> str | None:
    """Busca el ejecutable de Pandoc en el sistema o rutas estándar de Windows."""
    found = shutil.which("pandoc") or shutil.which("pandoc.exe")
    if found:
        return str(found)

    candidates = [
        Path(os.environ.get("LOCALAPPDATA", "")) / "Pandoc" / "pandoc.exe",
        Path(os.environ.get("PROGRAMFILES", "")) / "Pandoc" / "pandoc.exe",
        Path(os.environ.get("ProgramFiles(x86)", "")) / "Pandoc" / "pandoc.exe",
        Path.home() / "AppData" / "Local" / "Pandoc" / "pandoc.exe",
        Path(sys.executable).parent / "pandoc.exe",
        Path(sys.executable).parent / "Scripts" / "pandoc.exe",
    ]

    for p in candidates:
        if p and p.exists():
            return str(p)

    return None


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def _set_cell_margins(cell, top=110, bottom=110, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath("w:tblPr")
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)


def style_converted_docx(docx_path: Path) -> None:
    """Aplica la tipografía corporativa, jerarquía de títulos y estilos de tablas idénticos a LaTeX."""
    try:
        doc = docx.Document(str(docx_path))
    except Exception:
        return

    # Márgenes: 2.2 cm (0.86 in) igual que \geometry{margin=2.2cm} en LaTeX
    for section in doc.sections:
        section.top_margin = Inches(0.86)
        section.bottom_margin = Inches(0.86)
        section.left_margin = Inches(0.86)
        section.right_margin = Inches(0.86)

    # 1. Estilos de Párrafos y Encabezados
    for p in doc.paragraphs:
        style_name = p.style.name.lower()
        text_strip = p.text.strip()

        if "heading 1" in style_name:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.bold = True
                r.font.size = Pt(15)
                r.font.color.rgb = RGBColor(23, 107, 115)  # hbprimary
        elif "heading 2" in style_name:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.bold = True
                r.font.size = Pt(12.5)
                r.font.color.rgb = RGBColor(31, 157, 143)  # hbaccent
        elif "heading 3" in style_name:
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(15, 23, 42)
        elif "caption" in style_name or text_strip.startswith("Figura"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(9)
                r.font.italic = True
                r.font.color.rgb = RGBColor(100, 116, 139)
        else:
            p.paragraph_format.line_spacing = 1.15
            for r in p.runs:
                r.font.name = "Calibri"
                if not r.font.size:
                    r.font.size = Pt(10.5)
                if not r.font.color.rgb:
                    r.font.color.rgb = RGBColor(15, 23, 42)

    # 2. Estilos de Tablas idénticos a LaTeX
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, color="CBD5E1", sz="4", val="single")

        for r_idx, row in enumerate(table.rows):
            is_header = r_idx == 0
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

                if is_header:
                    _set_cell_bg(cell, "1F9D8F")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Calibri"
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(9.5)
                else:
                    if r_idx % 2 == 1:
                        _set_cell_bg(cell, "F8FAFC")
                    else:
                        _set_cell_bg(cell, "FFFFFF")
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(15, 23, 42)

    doc.save(str(docx_path))


def convert_latex_to_docx(tex_path: Path, output_docx_path: Path) -> bool:
    """Convierte un documento .tex a Word .docx directamente usando Pandoc y le aplica
    los estilos visuales, fuentes y tablas idénticos a la versión LaTeX.
    """
    tex_path = Path(tex_path).resolve()
    output_docx_path = Path(output_docx_path).resolve()
    pandoc_exe = find_pandoc()

    if not pandoc_exe or not tex_path.exists():
        return False

    # 1. Preprocesar código LaTeX para compatibilidad total con Pandoc
    raw_tex = tex_path.read_text(encoding="utf-8")
    clean_tex = re.sub(r"\\resizebox\{[^{}]*\}\{[^{}]*\}\{%\s*", "", raw_tex)
    clean_tex = re.sub(r"\\resizebox\{[^{}]*\}\{[^{}]*\}\{\s*", "", clean_tex)
    clean_tex = clean_tex.replace("\\end{tabular}%\n}", "\\end{tabular}")
    clean_tex = clean_tex.replace("\\end{tabular}%\n}}", "\\end{tabular}")
    clean_tex = clean_tex.replace("\\end{tabular}\n}", "\\end{tabular}")
    clean_tex = clean_tex.replace("\\end{tabular}\n}}", "\\end{tabular}")

    tmp_tex = tex_path.parent / "_pandoc_temp_input.tex"
    tmp_tex.write_text(clean_tex, encoding="utf-8")

    try:
        # Ejecutar Pandoc desde el directorio del informe para resolver las imágenes relativas (figuras/...)
        proc = subprocess.run(
            [
                pandoc_exe,
                tmp_tex.name,
                "-o",
                output_docx_path.name,
                "--from=latex",
                "--to=docx",
            ],
            cwd=str(tex_path.parent),
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )

        if proc.returncode == 0 and output_docx_path.exists():
            # 2. Aplicar tipografía, márgenes, colores y tablas estilizadas al Word generado
            style_converted_docx(output_docx_path)
            return True
        return False
    except Exception:
        return False
    finally:
        if tmp_tex.exists():
            try:
                tmp_tex.unlink()
            except Exception:
                pass
