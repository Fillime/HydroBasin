from __future__ import annotations

import math
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _format_table(table, col_widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        is_header = r_idx == 0
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell, top=110, bottom=110, left=160, right=160)
            if is_header:
                _set_cell_bg(cell, "1F9D8F")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(9.5)
            else:
                if r_idx % 2 == 1:
                    _set_cell_bg(cell, "F8FAFC")
                else:
                    _set_cell_bg(cell, "FFFFFF")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(15, 23, 42)
            if col_widths and c_idx < len(col_widths):
                cell.width = Inches(col_widths[c_idx])


def _n(val: object, decimals: int = 2) -> str:
    if val is None:
        return "N/D"
    try:
        v = float(val)
        if math.isnan(v) or math.isinf(v):
            return "N/D"
        return f"{v:.{decimals}f}"
    except (ValueError, TypeError):
        return str(val) if str(val).strip() else "N/D"


def generar_informe_docx(
    output_docx: Path,
    summary: dict,
    figures: dict[str, str],
    subbasins=None,
    loc: dict | None = None,
) -> Path:
    """Genera un informe técnico formal en formato Microsoft Word (.docx)."""
    loc = loc or {}
    output_docx = Path(output_docx).resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Configuración de márgenes (2.2 cm = 0.86 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.86)
        section.bottom_margin = Inches(0.86)
        section.left_margin = Inches(0.86)
        section.right_margin = Inches(0.86)

    site = summary.get("project_name") or "Cuenca Hidrográfica"
    client = summary.get("client") or "Particular"
    calc = summary.get("calculated_by") or "HydroBasin Studio"
    rev = summary.get("reviewed_by") or "Revisión Técnica"
    admin_label = summary.get("location_label") or "Colombia"
    outlet = summary.get("outlet_original") or {}
    total_area = float(summary.get("area_km2") or 1.0)

    # ============================== PORTADA ==============================
    p_top = doc.add_paragraph()
    r_top = p_top.add_run("HYDROBASIN STUDIO — INGENIERÍA HIDROLÓGICA")
    r_top.font.bold = True
    r_top.font.size = Pt(11)
    r_top.font.color.rgb = RGBColor(31, 157, 143)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    r_main = title_p.add_run("ESTUDIO HIDROLÓGICO Y MORFOMÉTRICO DE CUENCA\n")
    r_main.font.bold = True
    r_main.font.size = Pt(22)
    r_main.font.color.rgb = RGBColor(15, 23, 42)

    r_sub = title_p.add_run(f"{site}\n")
    r_sub.font.bold = True
    r_sub.font.size = Pt(15)
    r_sub.font.color.rgb = RGBColor(31, 157, 143)

    r_loc = title_p.add_run(f"{admin_label}\n\n")
    r_loc.font.size = Pt(11)
    r_loc.font.color.rgb = RGBColor(100, 116, 139)

    # Cuadro de Autoría y Revisión
    t_meta = doc.add_table(rows=3, cols=3)
    t_meta.rows[0].cells[0].paragraphs[0].add_run("ELABORÓ:")
    t_meta.rows[0].cells[1].paragraphs[0].add_run("REVISÓ:")
    t_meta.rows[0].cells[2].paragraphs[0].add_run("FECHA:")
    t_meta.rows[1].cells[0].paragraphs[0].add_run(calc)
    t_meta.rows[1].cells[1].paragraphs[0].add_run(rev)
    t_meta.rows[1].cells[2].paragraphs[0].add_run(datetime.now().strftime("%d/%m/%Y"))
    t_meta.rows[2].cells[0].paragraphs[0].add_run(f"CLIENTE: {client}")
    t_meta.rows[2].cells[1].paragraphs[0].add_run("ESTADO: Aprobado")
    t_meta.rows[2].cells[2].paragraphs[0].add_run("VERSIÓN: 1.0")
    _format_table(t_meta, [2.3, 2.3, 2.0])

    doc.add_page_break()

    # ============================== CAPÍTULOS ==============================
    def add_h1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor(23, 107, 115)
            r.font.bold = True
        return h

    def add_h2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.color.rgb = RGBColor(31, 157, 143)
            r.font.bold = True
        return h

    def add_fig(rel_path, caption):
        if not rel_path:
            return
        p_img = output_docx.parent / rel_path
        if p_img.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(str(p_img), width=Inches(5.6))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = cap.add_run(caption)
            r_cap.font.italic = True
            r_cap.font.size = Pt(8.5)
            r_cap.font.color.rgb = RGBColor(100, 116, 139)

    # 1. INTRODUCCIÓN
    add_h1("1. Introducción")
    doc.add_paragraph(
        f"El presente documento corresponde al informe técnico del estudio hidrológico y geomorfológico desarrollado para el proyecto {site}, localizado en la jurisdicción de {admin_label}. Comprende la delimitación de la divisoria de aguas, caracterización física de la cuenca, consulta de estaciones meteorológicas oficiales del IDEAM, cálculo de Polígonos de Thiessen, curvas IDF, estimación del Número de Curva SCS (CN), tiempos de concentración y caudales máximos de diseño."
    )

    # 2. OBJETIVOS Y ALCANCE
    add_h1("2. Objetivos y Alcance")
    add_h2("2.1 Objetivo General")
    doc.add_paragraph(
        f"Delimitar y caracterizar geomorfológica e hidrológicamente la cuenca aportante al punto de interés del proyecto {site}, determinando sus parámetros morfométricos, red de drenaje, tiempos de concentración y caudales de diseño."
    )
    add_h2("2.2 Objetivos Específicos")
    p_obj = doc.add_paragraph()
    p_obj.add_run("• Delimitar la cuenca hidrográfica y su red de drenaje a partir de modelos digitales de elevación corregidos.\n")
    p_obj.add_run("• Calcular los parámetros morfométricos de forma, relieve y red fluvial.\n")
    p_obj.add_run("• Subdividir la cuenca en subcuencas tributarias y caracterizar el cauce principal con su perfil altimétrico.\n")
    p_obj.add_run("• Consultar la red pluviométrica del IDEAM y ponderar su representatividad mediante Polígonos de Thiessen.\n")
    p_obj.add_run("• Formular las curvas IDF y determinar el Número de Curva ponderado (SCS-CN).\n")
    p_obj.add_run("• Estimar los tiempos de concentración y modelar los caudales de diseño para periodos de retorno de 2.33 a 100 años.")

    add_h2("2.3 Alcance del Estudio")
    doc.add_paragraph(
        "El alcance comprende el análisis morfométrico espacial, la modelación hidrometeorológica regional con datos oficiales del IDEAM y la formulación hidrológica mediante el método del Hidrograma Unitario del SCS y el Método Racional."
    )

    # 3. LOCALIZACIÓN Y UBICACIÓN GENERAL
    add_h1("3. Localización y Ubicación General")
    cota_ex = summary.get("main_channel_elevation_outlet_m") or summary.get("elevacion_min_m") or 0.0
    doc.add_paragraph(
        f"El área de estudio se ubica en la jurisdicción de {admin_label}. El punto de cierre o exutorio se localiza en las coordenadas geográficas {_n(outlet.get('y'),6)}° Latitud Norte, {_n(outlet.get('x'),6)}° Longitud Oeste (WGS84), a una cota de {_n(cota_ex, 1)} msnm."
    )
    add_fig(figures.get("location_satellite"), "Figura 1: Localización regional del área de estudio sobre imagen satelital.")

    # 4. INFORMACIÓN TOPOGRÁFICA Y ACONDICIONAMIENTO
    add_h1("4. Información Topográfica y Acondicionamiento del Terreno")
    doc.add_paragraph(
        f"La caracterización geomorfológica se fundamentó en el Modelo Digital de Elevación (DEM) {summary.get('dem_source') or 'DEM Satelital'}, con una resolución espacial de {_n((summary.get('metric_resolution_m') or [30])[0], 1)} m. Se aplicó un acondicionamiento hidrológico mediante el relleno de depresiones cerradas (sinks) para asegurar la continuidad física del escurrimiento."
    )
    add_fig(figures.get("dem"), "Figura 2: Contexto regional del Modelo Digital de Elevación (DEM).")
    add_fig(figures.get("hillshade"), "Figura 3: Relieve sombreado (Hillshade) del área de estudio.")
    add_fig(figures.get("flow_direction"), "Figura 4: Dirección de flujo D8 dentro de la cuenca.")
    add_fig(figures.get("accumulation"), "Figura 5: Acumulación de flujo dentro de la cuenca aportante.")

    # 5. CARACTERIZACIÓN MORFOMÉTRICA
    add_h1("5. Caracterización Morfométrica de la Cuenca")

    add_h2("5.1 Área, Perímetro y Longitud Axial")
    t_morf = doc.add_table(rows=14, cols=2)
    morf_data = [
        ("Parámetro Morfométrico", "Valor Obtenido"),
        ("Área de la Cuenca (A)", f"{_n(summary.get('area_km2'))} km²"),
        ("Perímetro de la Cuenca (P)", f"{_n(summary.get('perimetro_km'))} km"),
        ("Longitud Axial (La)", f"{_n(summary.get('longitud_axial_km'))} km"),
        ("Factor de Forma de Horton (Kf)", _n(summary.get('factor_forma'), 3)),
        ("Coeficiente de Compacidad de Gravelius (Kc)", _n(summary.get('coeficiente_compacidad'), 3)),
        ("Relación de Circularidad de Miller (Rc)", _n(summary.get('relacion_circularidad'), 3)),
        ("Densidad de Drenaje (Dd)", f"{_n(summary.get('densidad_drenaje_km_km2'), 3)} km/km²"),
        ("Orden Máximo de Corrientes (Strahler)", str(summary.get('strahler_max', 'N/D'))),
        ("Número de Subcuencas", str(summary.get('subbasin_count', 'N/D'))),
        ("Elevación Mínima (Exutorio)", f"{_n(summary.get('elevacion_min_m'))} msnm"),
        ("Elevación Máxima (Cabecera)", f"{_n(summary.get('elevacion_max_m'))} msnm"),
        ("Elevación Media de la Cuenca", f"{_n(summary.get('elevacion_media_m'))} msnm"),
        ("Relieve Total de la Cuenca (HT)", f"{_n(summary.get('relieve_cuenca_m'))} m"),
    ]
    for r_i, (k, v) in enumerate(morf_data):
        t_morf.rows[r_i].cells[0].paragraphs[0].add_run(k)
        t_morf.rows[r_i].cells[1].paragraphs[0].add_run(v)
    _format_table(t_morf, [4.2, 2.4])

    add_h2("5.2 Factor de Forma de Horton (Kf)")
    doc.add_paragraph(
        f"El factor de forma obtenido es Kf = {_n(summary.get('factor_forma'), 3)}, clasificándose como una cuenca {summary.get('clasificacion_factor_forma') or 'alargada'}, lo cual indica una baja susceptibilidad a crecientes súbitas simultáneas."
    )

    add_h2("5.3 Índice de Compacidad de Gravelius (Kc)")
    doc.add_paragraph(
        f"El índice de Gravelius arrojó un valor de Kc = {_n(summary.get('coeficiente_compacidad'), 3)}, confirmando una geometría alargada que favorece la disipación temporal de las ondas de avenida."
    )

    add_h2("5.4 Relación de Circularidad de Miller (Rc)")
    doc.add_paragraph(
        f"La relación de circularidad obtenida es Rc = {_n(summary.get('relacion_circularidad'), 3)}, ratificando la tendencia al amortiguamiento hidrodinámico."
    )

    add_h2("5.5 Densidad de Drenaje (Dd)")
    doc.add_paragraph(
        f"La densidad de drenaje calculada es Dd = {_n(summary.get('densidad_drenaje_km_km2'), 3)} km/km², reflejando la capacidad de evacuación de caudales superficiales y el grado de disección fluvial de la cuenca."
    )

    add_h2("5.6 Relieve y Elevaciones")
    doc.add_paragraph(
        f"El análisis altimétrico revela una cota mínima de {_n(summary.get('elevacion_min_m'))} msnm, máxima de {_n(summary.get('elevacion_max_m'))} msnm y media de {_n(summary.get('elevacion_media_m'))} msnm, representando un relieve total de HT = {_n(summary.get('relieve_cuenca_m'))} m."
    )

    add_h2("5.7 Orden de Strahler")
    doc.add_paragraph(
        f"La red fluvial alcanza un orden de corrientes de Strahler máximo de {summary.get('strahler_max', 'N/D')}."
    )
    add_fig(figures.get("strahler"), "Figura 6: Jerarquía de la red fluvial según el Orden de Corrientes de Strahler.")

    add_h2("5.8 Subcuencas y Red de Drenaje")
    doc.add_paragraph(
        f"Se identificaron y delimitaron {summary.get('subbasin_count', 'N/D')} subcuencas hidrológicas internas."
    )
    add_fig(figures.get("watershed"), "Figura 7: Cuenca delimitada, red de drenaje y cauce principal.")
    add_fig(figures.get("subbasins"), "Figura 8: Subdivisión en subcuencas hidrológicas.")

    add_h2("5.9 Cauce Principal y Perfil Longitudinal")
    desnivel = (summary.get("main_channel_elevation_source_m") or 0) - (summary.get("main_channel_elevation_outlet_m") or 0)
    t_ch = doc.add_table(rows=6, cols=2)
    ch_data = [
        ("Parámetro del Cauce", "Valor"),
        ("Longitud del Cauce Principal (L)", f"{_n(summary.get('main_channel_length_km'))} km"),
        ("Elevación en Cabecera", f"{_n(summary.get('main_channel_elevation_source_m'))} msnm"),
        ("Elevación en Exutorio", f"{_n(summary.get('main_channel_elevation_outlet_m'))} msnm"),
        ("Desnivel Topográfico (ΔH)", f"{_n(desnivel)} m"),
        ("Pendiente Media del Cauce (S)", f"{_n(summary.get('main_channel_slope_percent'), 3)}%"),
    ]
    for r_i, (k, v) in enumerate(ch_data):
        t_ch.rows[r_i].cells[0].paragraphs[0].add_run(k)
        t_ch.rows[r_i].cells[1].paragraphs[0].add_run(v)
    _format_table(t_ch, [4.2, 2.4])
    add_fig(figures.get("profile"), "Figura 9: Perfil longitudinal altimétrico del cauce principal.")

    add_h2("5.10 Cartografía Morfométrica Complementaria")
    doc.add_paragraph(
        "La cartografía morfométrica y los vectores GeoPackage generados constituyen la base geométrica para el dimensionamiento hidráulico y zonificación ambiental."
    )

    # 6. CARACTERIZACIÓN HIDROLÓGICA
    add_h1("6. Caracterización Hidrológica")

    add_h2("6.1 Estaciones IDEAM")
    stations = summary.get("ideam_stations") or []
    t_st = doc.add_table(rows=len(stations[:8]) + 1, cols=7)
    st_headers = ["Código", "Nombre", "Categoría", "Altitud", "Lat / Lon", "Municipio", "Distancia"]
    for c_i, h in enumerate(st_headers):
        t_st.rows[0].cells[c_i].paragraphs[0].add_run(h)
    for r_i, s in enumerate(stations[:8], start=1):
        muni = s.get("municipio") or admin_label.split(",")[0]
        t_st.rows[r_i].cells[0].paragraphs[0].add_run(str(s.get("codigo", "")))
        t_st.rows[r_i].cells[1].paragraphs[0].add_run(str(s.get("nombre", "")))
        t_st.rows[r_i].cells[2].paragraphs[0].add_run(str(s.get("categoria", "")))
        t_st.rows[r_i].cells[3].paragraphs[0].add_run(f"{_n(s.get('altitud'), 0)} m")
        t_st.rows[r_i].cells[4].paragraphs[0].add_run(f"{_n(s.get('latitud'), 4)}° / {_n(s.get('longitud'), 4)}°")
        t_st.rows[r_i].cells[5].paragraphs[0].add_run(muni)
        t_st.rows[r_i].cells[6].paragraphs[0].add_run(f"{_n(s.get('distancia_km'), 1)} km")
    _format_table(t_st)
    add_fig(figures.get("stations_map"), "Figura 10: Ubicación espacial de las estaciones meteorológicas del IDEAM.")

    add_h2("6.2 Polígonos de Thiessen")
    th_weights = summary.get("thiessen_weights") or []
    t_th = doc.add_table(rows=len(th_weights) + 1, cols=4)
    th_headers = ["Código", "Estación Meteorológica", "Área de Influencia (km²)", "% Área Cuenca"]
    for c_i, h in enumerate(th_headers):
        t_th.rows[0].cells[c_i].paragraphs[0].add_run(h)
    for r_i, th in enumerate(th_weights, start=1):
        t_th.rows[r_i].cells[0].paragraphs[0].add_run(str(th.get("codigo", "")))
        t_th.rows[r_i].cells[1].paragraphs[0].add_run(str(th.get("nombre", "")))
        t_th.rows[r_i].cells[2].paragraphs[0].add_run(_n(th.get("area_km2")))
        t_th.rows[r_i].cells[3].paragraphs[0].add_run(f"{_n(th.get('porcentaje'), 1)}%")
    _format_table(t_th, [1.4, 2.4, 1.6, 1.2])
    add_fig(figures.get("thiessen_map"), "Figura 11: Polígonos de Thiessen y áreas de influencia pluviométrica.")

    add_h2("6.3 Curvas IDF")
    doc.add_paragraph(
        "A partir de las ecuaciones regionales del IDEAM se construyeron las Curvas Intensidad-Duración-Frecuencia (IDF) para periodos de retorno de 2.33 a 100 años."
    )
    add_fig(figures.get("idf_curves"), "Figura 12: Curvas Intensidad–Duración–Frecuencia (IDF) calculadas.")

    add_h2("6.4 Número de Curva SCS-CN")
    doc.add_paragraph(
        f"El Número de Curva ponderado obtenido para la cuenca es CN = {_n(summary.get('cn_weighted'), 1)}, con una retención potencial máxima de humedad S = {_n(summary.get('curve_number', {}).get('s_retention_mm'), 1)} mm y abstracción inicial Ia = {_n(summary.get('curve_number', {}).get('ia_abstraction_mm'), 1)} mm."
    )
    add_fig(figures.get("curve_number"), "Figura 13: Distribución de coberturas y Número de Curva ponderado.")

    add_h2("6.5 Tiempo de Concentración")
    tc_avg_h = float(summary.get("tc_promedio_h") or 1.0)
    tc_avg_min = tc_avg_h * 60.0
    doc.add_paragraph(
        f"El tiempo de concentración medio adoptado es Tc = {_n(tc_avg_min, 1)} minutos ({_n(tc_avg_h, 2)} horas), fundamentado en el promedio de formulaciones hidrológicas (Kirpich, Témez, Giandotti, Johnstone & Cross, Chow)."
    )

    add_h2("6.6 Modelación Hidrológica y Caudales Máximos")
    peak_flows = summary.get("peak_discharges") or []
    t_q = doc.add_table(rows=len(peak_flows) + 1, cols=7)
    q_headers = ["Tr (años)", "I (mm/h)", "Ptotal (mm)", "Pefectiva (mm)", "Q Racional (m³/s)", "Q SCS (m³/s)", "Q Diseño (m³/s)"]
    for c_i, h in enumerate(q_headers):
        t_q.rows[0].cells[c_i].paragraphs[0].add_run(h)
    for r_i, q in enumerate(peak_flows, start=1):
        t_q.rows[r_i].cells[0].paragraphs[0].add_run(f"Tr = {q['tr_anos']} a")
        t_q.rows[r_i].cells[1].paragraphs[0].add_run(_n(q['intensidad_mm_h'], 1))
        t_q.rows[r_i].cells[2].paragraphs[0].add_run(_n(q['precipitacion_total_mm'], 1))
        t_q.rows[r_i].cells[3].paragraphs[0].add_run(_n(q['precipitacion_efectiva_mm'], 1))
        t_q.rows[r_i].cells[4].paragraphs[0].add_run(_n(q['caudal_racional_m3_s']))
        t_q.rows[r_i].cells[5].paragraphs[0].add_run(_n(q['caudal_scs_m3_s']))
        t_q.rows[r_i].cells[6].paragraphs[0].add_run(_n(q['caudal_diseno_m3_s']))
    _format_table(t_q)
    add_fig(figures.get("hydrographs"), "Figura 14: Hidrogramas de diseño para diferentes periodos de retorno.")

    # 7. ANÁLISIS INTEGRADO DE RESULTADOS
    add_h1("7. Análisis Integrado de Resultados")
    doc.add_paragraph(
        f"La cuenca presenta un factor de forma Kf = {_n(summary.get('factor_forma'), 3)} e índice de compacidad Kc = {_n(summary.get('coeficiente_compacidad'), 3)}, lo que se traduce en una respuesta hidrodinámica amortiguada. Para el periodo de retorno de diseño Tr = 100 años, el caudal pico estimado de {_n(peak_flows[5]['caudal_diseno_m3_s'] if len(peak_flows)>5 else 0)} m³/s representa la condición de solicitación extrema recomendada para obras de protección hidráulica."
    )

    # 8. CONCLUSIONES Y RECOMENDACIONES
    add_h1("8. Conclusiones y Recomendaciones")
    p_conc = doc.add_paragraph()
    p_conc.add_run(f"• Se delimitó exitosamente la cuenca aportante para {site}, con un área de {_n(summary.get('area_km2'))} km² y un perímetro de {_n(summary.get('perimetro_km'))} km.\n")
    p_conc.add_run(f"• El tiempo de concentración medio adoptado es de {_n(tc_avg_min, 1)} minutos ({_n(tc_avg_h, 2)} h).\n")
    p_conc.add_run(f"• El Número de Curva ponderado estimado es CN = {_n(summary.get('cn_weighted'), 1)}.\n")
    p_conc.add_run(f"• Los caudales máximos de diseño modelados para Tr = 25, 50 y 100 años son {_n(peak_flows[3]['caudal_diseno_m3_s'] if len(peak_flows)>3 else 0)} m³/s, {_n(peak_flows[4]['caudal_diseno_m3_s'] if len(peak_flows)>4 else 0)} m³/s y {_n(peak_flows[5]['caudal_diseno_m3_s'] if len(peak_flows)>5 else 0)} m³/s respectivamente.")

    # 9. LIMITACIONES TÉCNICAS
    add_h1("9. Limitaciones Técnicas")
    doc.add_paragraph(
        "Los resultados derivan del procesamiento topográfico de modelos de elevación satelitales y formulaciones hidrológicas sintéticas. El presente estudio no reemplaza levantamientos batimétricos directos en el cauce ni la verificación estructural en campo de las obras de paso existentes."
    )

    # 10. ANEXOS IDEAM
    doc.add_page_break()
    add_h1("10. Anexos IDEAM — Registro y Tratamiento de Información Pluviométrica")
    doc.add_paragraph(
        "A continuación se presenta el registro de estaciones del IDEAM consultadas por HydroBasin junto con las metodologías técnicas estandarizadas para la completación y homogeneización de series."
    )

    add_h2("10.1 Registro Oficial de Estaciones IDEAM")
    t_anx = doc.add_table(rows=len(stations) + 1, cols=8)
    anx_headers = ["Código", "Nombre", "Categoría", "Altitud", "Latitud", "Longitud", "Municipio", "Estado"]
    for c_i, h in enumerate(anx_headers):
        t_anx.rows[0].cells[c_i].paragraphs[0].add_run(h)
    for r_i, s in enumerate(stations, start=1):
        muni = s.get("municipio") or admin_label.split(",")[0]
        t_anx.rows[r_i].cells[0].paragraphs[0].add_run(str(s.get("codigo", "N/D")))
        t_anx.rows[r_i].cells[1].paragraphs[0].add_run(str(s.get("nombre", "N/D")))
        t_anx.rows[r_i].cells[2].paragraphs[0].add_run(str(s.get("categoria", "N/D")))
        t_anx.rows[r_i].cells[3].paragraphs[0].add_run(f"{_n(s.get('altitud'), 0)} m")
        t_anx.rows[r_i].cells[4].paragraphs[0].add_run(f"{_n(s.get('latitud'), 4)}°")
        t_anx.rows[r_i].cells[5].paragraphs[0].add_run(f"{_n(s.get('longitud'), 4)}°")
        t_anx.rows[r_i].cells[6].paragraphs[0].add_run(muni)
        t_anx.rows[r_i].cells[7].paragraphs[0].add_run(str(s.get("estado", "Activa")))
    _format_table(t_anx)

    add_h2("10.2 Metodologías de Completación y Homogeneización")
    p_meth = doc.add_paragraph()
    p_meth.add_run("1. Método de Proporciones Normales (Normal Ratio Method): Aplicado cuando la precipitación media anual de las estaciones vecinas difiere en más de un 10% respecto a la estación objetivo.\n")
    p_meth.add_run("2. Regresión Lineal y Correlación Cruzada: Aplicado en estaciones con R² >= 0.75 y regímenes climáticos consistentes.\n")
    p_meth.add_run("3. Interpolación Espacial por Inverso de la Distancia Ponderada (IDW): Empleado para estimar campos espaciales continuos de precipitación.")

    doc.save(str(output_docx))
    return output_docx
